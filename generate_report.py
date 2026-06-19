# -*- coding: utf-8 -*-
"""Generate EduCenter Final Project Report (DOCX)."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "FINAL_PROJECT_REPORT_EduCenter.docx")


def set_doc_defaults(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def add_center_title(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def add_screenshot_marker(doc, title, file_path, note=""):
    p = doc.add_paragraph()
    run = p.add_run(f"[INSERT SCREENSHOT HERE — {title}]")
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    run.font.name = "Times New Roman"
    add_para(doc, f"Source file: {file_path}", italic=True)
    if note:
        add_para(doc, note, italic=True)
    doc.add_paragraph()


def add_figure_placeholder(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[INSERT IMAGE HERE — {caption}]")
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    run.font.name = "Times New Roman"
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    doc.add_paragraph()


def add_table_schema(doc, title, rows):
    add_para(doc, title, bold=True)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Attribute", "Data Type", "Width/Constraint", "Meaning"]):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def build_report():
    doc = Document()
    set_doc_defaults(doc)

    # ===== COVER PAGE =====
    for _ in range(3):
        doc.add_paragraph()
    add_center_title(doc, "VIETNAM NATIONAL UNIVERSITY, HANOI", 13)
    add_center_title(doc, "INTERNATIONAL SCHOOL", 13)
    add_center_title(doc, "INS3064", 13)
    add_center_title(doc, "MULTIMEDIA DESIGN AND WEB DEVELOPMENT", 13)
    doc.add_paragraph()
    add_center_title(doc, "FINAL PROJECT REPORT", 16)
    add_center_title(doc, "EDUCATION CENTER MANAGEMENT SYSTEM (EduCenter)", 14)
    doc.add_paragraph()
    add_center_title(doc, "Group members", 12, False)
    add_center_title(doc, "[Student Name 1] - [Student ID 1]", 12, False)
    add_center_title(doc, "[Student Name 2] - [Student ID 2]", 12, False)
    add_center_title(doc, "[Student Name 3] - [Student ID 3]", 12, False)
    add_center_title(doc, "[Student Name 4] - [Student ID 4]", 12, False)
    doc.add_paragraph()
    add_center_title(doc, "SUPERVISOR: PhD. Phạm Đức Thọ", 12, False)
    doc.add_paragraph()
    add_center_title(doc, "Hanoi – Year 2026", 12, False)
    doc.add_page_break()

    # ===== ACKNOWLEDGEMENTS =====
    add_heading(doc, "ACKNOWLEDGEMENTS", 1)
    add_para(doc, (
        "We wish to extend our most profound gratitude to our supervising lecturer, whose scholarly "
        "guidance and insightful mentorship have been the cornerstone of this project. Their unwavering "
        "support and constructive critiques provided us with the intellectual clarity necessary to navigate "
        "the complexities of this course and bring this final endeavor to a successful realization."
    ))
    add_para(doc, (
        "Our heartfelt thanks also go to every member of the project team. This achievement is a testament "
        "to our collective commitment, seamless synergy, and the relentless pursuit of excellence that defined "
        "our collaboration."
    ))
    add_para(doc, (
        "Finally, we express our sincere appreciation to the Faculty and the University for the exceptional "
        "academic environment and the comprehensive resources provided."
    ))
    doc.add_page_break()

    # ===== STUDENT DECLARATION =====
    add_heading(doc, "STUDENT DECLARATION", 1)
    add_para(doc, (
        "I certify that the assignment submission is entirely my own work and I fully understand the "
        "consequences of plagiarism. I declare that the work submitted for assessment has been carried out "
        "without assistance other than that which is acceptable according to the rules of the specification. "
        "I certify I have clearly referenced any sources and any artificial intelligence (AI) tools used "
        "in the work. I understand that making a false declaration is a form of malpractice."
    ))
    add_para(doc, "Student signature(s): _________________________    Date: ______________")
    add_para(doc, "[Student Name 1]    [Student Name 2]")
    add_para(doc, "[Student Name 3]    [Student Name 4]")
    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    add_heading(doc, "Table of Contents", 1)
    toc = [
        "CHAPTER 1: INTRODUCTION ........................................................ 7",
        "  1.1 Background And Concept / Introduction .................................... 7",
        "  1.2 Problem Statement ........................................................ 8",
        "  1.3 Project Question ......................................................... 8",
        "  1.4 Project Objectives ....................................................... 9",
        "  1.5 Scope Of The Project ..................................................... 9",
        "  1.6 Real-world Problem and Why It Matters .................................... 10",
        "CHAPTER 2: SYSTEM DESIGN ....................................................... 11",
        "  2.1 User Stories ............................................................. 11",
        "  2.2 Site Map ................................................................. 13",
        "  2.3 Entity Relationship Diagram .............................................. 14",
        "CHAPTER 3: IMPLEMENTATION .................................................... 18",
        "  3.1 Sample Source Code ....................................................... 18",
        "  3.2 Images of Final Application .............................................. 24",
        "  3.3 GitHub Repository Evidences .............................................. 27",
        "CHAPTER 4: CONCLUSION .......................................................... 28",
        "  4.1 What Went Well ........................................................... 28",
        "  4.2 What Did Not Go Well ..................................................... 29",
        "  4.3 Lessons Learned and Further Improvements ................................. 29",
    ]
    for line in toc:
        add_para(doc, line)
    doc.add_page_break()

    # ===== LIST OF FIGURES =====
    add_heading(doc, "List of Figures", 1)
    figures = [
        "Figure 2.1: Site map of EduCenter",
        "Figure 2.2: Entity-Relationship Diagram",
        "Figure 3.1: Admin Dashboard",
        "Figure 3.2: Login page",
        "Figure 3.3: Course management page (Admin)",
        "Figure 3.4: Class and enrollment management",
        "Figure 3.5: Teacher schedule page",
        "Figure 3.6: Student schedule page",
        "Figure 3.7: Student submission (PDF upload) page",
        "Figure 3.8: Teacher grading page",
        "Figure 3.9: GitHub commit history (Member 1)",
        "Figure 3.10: GitHub commit history (Member 2)",
        "Figure 3.11: GitHub commit history (Member 3)",
    ]
    for f in figures:
        add_para(doc, f)
    doc.add_page_break()

    # ===== LIST OF TABLES =====
    add_heading(doc, "List of Tables", 1)
    tables = [
        "Table 2.1: users table",
        "Table 2.2: roles and user_roles tables",
        "Table 2.3: courses table",
        "Table 2.4: classes and enrollments tables",
        "Table 2.5: schedules and teacher_assignments tables",
        "Table 2.6: submissions and grades tables",
    ]
    for t in tables:
        add_para(doc, t)
    doc.add_page_break()

    # ============================================================
    # CHAPTER 1
    # ============================================================
    add_heading(doc, "CHAPTER 1: INTRODUCTION", 1)

    add_heading(doc, "1.1 Background And Concept / Introduction", 2)
    add_para(doc, (
        "In the context of digital transformation in education, private training centers and language "
        "schools increasingly need integrated software to manage courses, classes, teachers, students, "
        "schedules, tuition payments, attendance, and academic performance. Traditional management using "
        "spreadsheets, paper forms, and disconnected tools leads to data inconsistency, scheduling conflicts, "
        "delayed payment processing, and poor visibility for administrators, teachers, and students."
    ))
    add_para(doc, (
        "This project addresses that gap by developing EduCenter — a web-based Education Center Management "
        "System built with Native PHP and MySQL. The system provides a centralized portal for three user roles: "
        "Administrator, Teacher, and Student. It supports the full academic lifecycle from course planning and "
        "class creation to enrollment, payment approval, schedule generation, attendance tracking, assignment "
        "submission, grading, and reporting."
    ))
    add_para(doc, (
        "The application follows a custom MVC (Model-View-Controller) architecture. The web portal entry point is "
        "project_root/public/web.php, routes are defined in project_root/app/web_routes.php, business logic is "
        "handled by controllers and service models, and the user interface is rendered through PHP view templates "
        "under project_root/app/view/. Session-based authentication with role-based access control (RBAC) ensures "
        "that each user only accesses features appropriate to their role."
    ))

    add_heading(doc, "1.2 Problem Statement", 2)
    add_para(doc, (
        "Many education centers still rely on manual processes that create the following problems:"
    ))
    add_bullets(doc, [
        "Scheduling conflicts when students or teachers are assigned to overlapping time slots across multiple classes.",
        "Delayed tuition confirmation, causing students to remain without an active schedule even after payment.",
        "Difficulty tracking attendance, leave requests, makeup sessions, and payroll-related teaching hours.",
        "No centralized system for students to submit assignments (PDF) and for teachers to grade and provide feedback.",
        "Administrators lack real-time dashboards for revenue, enrollment status, and operational reports.",
        "Data is scattered across spreadsheets, making audit trails and decision-making inefficient.",
    ])

    add_heading(doc, "1.3 Project Question", 2)
    add_para(doc, (
        "How can a web application built with PHP and MySQL effectively automate the management of an education "
        "center — including course/class administration, teacher assignment, student enrollment and payment, "
        "scheduling with conflict detection, attendance, assignment submission, and grading — while providing "
        "role-based portals for administrators, teachers, and students?"
    ))

    add_heading(doc, "1.4 Project Objectives", 2)
    add_bullets(doc, [
        "Design and implement a secure, role-based web portal for Admin, Teacher, and Student users.",
        "Build CRUD modules for courses, classes, classrooms, semesters, students, teachers, and enrollments.",
        "Implement tuition payment workflow with schedule conflict checks before activation.",
        "Generate and manage class schedules (regular, exam, makeup) with conflict detection for teachers and students.",
        "Enable teachers to mark attendance, grade PDF submissions, and evaluate student performance.",
        "Enable students to view schedules, submit assignments, track results, and complete teacher surveys.",
        "Provide admin reports, audit logs, notifications, and export functionality (CSV).",
        "Design a normalized relational database schema with at least 20 related tables.",
    ])

    add_heading(doc, "1.5 Scope Of The Project", 2)
    add_para(doc, "In scope:", bold=True)
    add_bullets(doc, [
        "Web portal (session authentication) for Admin, Teacher, and Student roles.",
        "MySQL database with migration script (migrate.php) and seed data.",
        "Core modules: courses, classes, enrollments, payments, schedules, assignments, attendance, submissions, grades.",
        "Leave/makeup request workflow with admin approval.",
        "Server-rendered PHP views with Bootstrap Icons and custom CSS.",
        "REST JSON API (secondary entry via public/index.php) with JWT authentication.",
    ])
    add_para(doc, "Out of scope:", bold=True)
    add_bullets(doc, [
        "Mobile native applications (iOS/Android).",
        "Online payment gateway integration (payments are manually confirmed by admin).",
        "Video conferencing or live streaming integration.",
        "Multi-branch / multi-center enterprise deployment.",
    ])

    add_heading(doc, "1.6 Describe the Real-World Problem and Why It Matters", 2)
    add_para(doc, (
        "A real-world education center must coordinate hundreds of operational decisions every week: which "
        "teacher teaches which student, whether a classroom is available, whether a student has paid tuition, "
        "whether two classes overlap in time, and whether assignments have been graded on time. When these "
        "processes are manual, errors are inevitable — for example, a student may be enrolled in two classes "
        "with conflicting schedules, or a teacher may be double-booked."
    ))
    add_para(doc, (
        "EduCenter matters because it digitizes these workflows into a single system. When an administrator "
        "confirms a tuition payment, the system can automatically check schedule conflicts and activate the "
        "student's class schedule. Teachers can mark attendance linked to specific session dates, and students "
        "can upload PDF assignments and view grades online. This reduces administrative workload, improves data "
        "accuracy, and enhances transparency for all stakeholders — directly supporting the center's reputation "
        "and operational efficiency."
    ))
    doc.add_page_break()

    # ============================================================
    # CHAPTER 2
    # ============================================================
    add_heading(doc, "CHAPTER 2: SYSTEM DESIGN", 1)

    add_heading(doc, "2.1 User Stories", 2)
    add_para(doc, (
        "The following user stories were produced using the standard User Story template: "
        "\"As a [role], I want [feature], so that [benefit].\""
    ))

    add_para(doc, "Administrator Stories", bold=True)
    stories_admin = [
        "As an Admin, I want to manage courses and classes, so that I can organize the center's academic offerings.",
        "As an Admin, I want to enroll students into classes, so that each student is registered for the correct program.",
        "As an Admin, I want to confirm tuition payments, so that students become ACTIVE and receive their schedules.",
        "As an Admin, I want the system to detect schedule conflicts before confirming payment, so that no student is assigned to overlapping classes.",
        "As an Admin, I want to assign teachers to classes and students, so that teaching responsibilities are clearly defined.",
        "As an Admin, I want to view a dashboard with revenue and enrollment statistics, so that I can monitor business performance.",
        "As an Admin, I want to approve or reject leave/makeup requests, so that schedule changes are controlled.",
        "As an Admin, I want to export reports and audit logs, so that I can review system activity.",
    ]
    add_bullets(doc, stories_admin)

    add_para(doc, "Teacher Stories", bold=True)
    stories_teacher = [
        "As a Teacher, I want to view my weekly teaching schedule, so that I know when and where to teach.",
        "As a Teacher, I want to navigate between weeks (Previous/Next), so that I can plan ahead easily.",
        "As a Teacher, I want to mark student attendance for each session, so that attendance records are accurate.",
        "As a Teacher, I want to grade student PDF submissions, so that students receive timely feedback.",
        "As a Teacher, I want to submit leave/makeup requests, so that schedule changes are formally recorded.",
        "As a Teacher, I want to receive notifications when students submit assignments, so that I can grade promptly.",
    ]
    add_bullets(doc, stories_teacher)

    add_para(doc, "Student Stories", bold=True)
    stories_student = [
        "As a Student, I want to view my class schedule by week, so that I know when to attend classes.",
        "As a Student, I want to navigate between weeks (Previous/Next), so that I can check future sessions.",
        "As a Student, I want to view my attendance status, so that I know whether I was marked present or absent.",
        "As a Student, I want to upload PDF assignments, so that I can submit homework online.",
        "As a Student, I want to view my grades and teacher comments, so that I can track my academic progress.",
        "As a Student, I want to compare my scores with classmates, so that I understand my relative performance.",
        "As a Student, I want to submit leave/makeup requests, so that absences are officially recorded.",
        "As a Student, I want to complete a teacher survey after the course ends, so that I can provide feedback.",
    ]
    add_bullets(doc, stories_student)

    add_heading(doc, "2.2 Site Map", 2)
    add_para(doc, (
        "The EduCenter site map is organized by user role. After login, each user is redirected to their "
        "role-specific dashboard. The structure below shows the main navigation paths."
    ))
    sitemap = """
EduCenter Portal
├── Login / Logout
│   └── project_root/app/view/auth/login.php
│
├── Admin Portal (/admin/*)
│   ├── Dashboard                    → /admin/dashboard
│   ├── Notifications                → /admin/notifications
│   ├── Courses (CRUD)               → /admin/courses
│   ├── Classes (CRUD)               → /admin/classes
│   ├── Classrooms (CRUD)            → /admin/classrooms
│   ├── Semesters (CRUD)             → /admin/semesters
│   ├── Class Plans (CRUD)           → /admin/class-plans
│   ├── Students (CRUD)              → /admin/students
│   ├── Teachers (CRUD)              → /admin/teachers
│   ├── Enrollments (CRUD)           → /admin/enrollments
│   ├── Payments                     → /admin/payments
│   ├── Payrolls                     → /admin/payrolls
│   ├── Schedules                    → /admin/schedules
│   ├── Teacher Assignments          → /admin/assignments
│   ├── Teaching Schedule            → /admin/teaching-schedule
│   ├── Leave / Makeup Approval      → /admin/leave-requests
│   ├── Attendance Report            → /admin/attendance-report
│   ├── Reports & Export             → /admin/reports
│   └── Audit Logs                   → /admin/audit-logs
│
├── Teacher Portal (/teacher/*)
│   ├── Dashboard                    → /teacher/dashboard
│   ├── Notifications                → /teacher/notifications
│   ├── Teaching Schedule            → /teacher/schedule
│   ├── Attendance                   → /teacher/attendance
│   ├── Leave / Makeup               → /teacher/leave
│   └── Grading                      → /teacher/grading
│
└── Student Portal (/student/*)
    ├── Dashboard                    → /student/dashboard
    ├── Notifications                → /student/notifications
    ├── Class Schedule               → /student/schedule
    ├── Attendance                   → /student/attendance
    ├── Leave / Makeup               → /student/leave
    ├── Submissions (PDF)            → /student/submissions
    ├── Results                      → /student/results
    ├── Score Comparison             → /student/compare
    └── Teacher Survey               → /student/survey
"""
    p = doc.add_paragraph(sitemap)
    for run in p.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    add_figure_placeholder(doc, "Figure 2.1: Site map of EduCenter (draw in draw.io / Lucidchart and insert here)")

    add_heading(doc, "2.3 Entity Relationship Diagram", 2)
    add_para(doc, (
        "The database uses a relational model with 24 tables. The core entities are users, roles, teachers, "
        "students, courses, classes, enrollments, schedules, and teacher_assignments. Key relationships include:"
    ))
    add_bullets(doc, [
        "users ↔ user_roles ↔ roles (M:N) — role-based access control.",
        "users → teachers / students (1:1) — profile extension.",
        "courses → classes (1:N) — each class belongs to one course.",
        "classes ↔ students via enrollments (M:N) — with payment_status and status.",
        "classes → schedules (1:N) — weekly or specific-date sessions.",
        "teachers ↔ classes ↔ students via teacher_assignments (M:N with attributes).",
        "students → submissions → grades (1:N chain) — assignment workflow.",
        "students → attendance (1:N) — per class per date.",
        "classes → surveys → survey_responses (1:N) — end-of-course feedback.",
    ])

    add_figure_placeholder(doc, "Figure 2.2: Entity-Relationship Diagram (draw in MySQL Workbench / draw.io using finaldemo_web.sql and insert here)")

    add_para(doc, "Physical database design (selected tables):", bold=True)

    add_table_schema(doc, "Table 2.1. users Table", [
        ("id", "INT", "PK, AUTO_INCREMENT", "Unique user identifier"),
        ("full_name", "VARCHAR(150)", "NOT NULL", "Full name"),
        ("email", "VARCHAR(150)", "UNIQUE, NOT NULL", "Login email"),
        ("password_hash", "VARCHAR(255)", "NOT NULL", "Bcrypt hashed password"),
        ("phone", "VARCHAR(20)", "NULL", "Phone number"),
        ("status", "ENUM", "ACTIVE/INACTIVE", "Account status"),
        ("created_at", "TIMESTAMP", "DEFAULT NOW", "Registration date"),
    ])

    add_table_schema(doc, "Table 2.2. roles and user_roles Tables", [
        ("roles.id", "INT", "PK", "Role identifier"),
        ("roles.role_name", "VARCHAR(60)", "UNIQUE", "ADMIN, TEACHER, STUDENT"),
        ("user_roles.user_id", "INT", "FK → users", "User reference"),
        ("user_roles.role_id", "INT", "FK → roles", "Role reference"),
    ])

    add_table_schema(doc, "Table 2.3. courses Table", [
        ("id", "INT", "PK", "Course identifier"),
        ("course_code", "VARCHAR(30)", "UNIQUE", "e.g. JS101, PY501"),
        ("course_name", "VARCHAR(200)", "NOT NULL", "Course title"),
        ("total_sessions", "INT", "DEFAULT 20", "Number of sessions"),
        ("day_primary", "TINYINT", "DEFAULT 1", "Primary weekday (Mon=1)"),
        ("day_secondary", "TINYINT", "DEFAULT 4", "Secondary weekday"),
        ("tuition_fee", "DECIMAL(12,2)", "NOT NULL", "Course fee"),
        ("status", "ENUM", "ACTIVE/INACTIVE", "Course status"),
    ])

    add_table_schema(doc, "Table 2.4. classes and enrollments Tables", [
        ("classes.id", "INT", "PK", "Class identifier"),
        ("classes.class_code", "VARCHAR(40)", "UNIQUE", "e.g. JS101-A"),
        ("classes.course_id", "INT", "FK → courses", "Parent course"),
        ("classes.teacher_id", "INT", "FK → teachers", "Assigned teacher"),
        ("enrollments.student_id", "INT", "FK → students", "Enrolled student"),
        ("enrollments.class_id", "INT", "FK → classes", "Target class"),
        ("enrollments.payment_status", "ENUM", "UNPAID/PAID", "Tuition status"),
        ("enrollments.status", "ENUM", "PENDING/ACTIVE", "Enrollment status"),
    ])

    add_table_schema(doc, "Table 2.5. schedules and teacher_assignments Tables", [
        ("schedules.class_id", "INT", "FK → classes", "Class reference"),
        ("schedules.student_id", "INT", "FK → students, NULL", "Per-student schedule"),
        ("schedules.day_of_week", "TINYINT", "0-6", "Day of week"),
        ("schedules.schedule_type", "ENUM", "REGULAR/EXAM/MAKEUP", "Session type"),
        ("teacher_assignments.teacher_id", "INT", "FK → teachers", "Assigned teacher"),
        ("teacher_assignments.class_id", "INT", "FK → classes", "Target class"),
        ("teacher_assignments.student_id", "INT", "FK → students", "Target student"),
    ])

    add_table_schema(doc, "Table 2.6. submissions and grades Tables", [
        ("submissions.student_id", "INT", "FK → students", "Submitting student"),
        ("submissions.class_id", "INT", "FK → classes", "Target class"),
        ("submissions.type", "ENUM", "ASSIGNMENT/MIDTERM/FINAL", "Submission type"),
        ("submissions.file_path", "VARCHAR(500)", "NULL", "Uploaded PDF path"),
        ("submissions.status", "ENUM", "PENDING/GRADED", "Grading status"),
        ("grades.submission_id", "INT", "FK → submissions", "Graded submission"),
        ("grades.score", "DECIMAL(5,2)", "NULL", "Score out of 10"),
        ("grades.comment", "TEXT", "NULL", "Teacher feedback"),
    ])
    doc.add_page_break()

    # ============================================================
    # CHAPTER 3
    # ============================================================
    add_heading(doc, "CHAPTER 3: IMPLEMENTATION", 1)

    add_heading(doc, "3.1 Sample Source Code", 2)
    add_para(doc, (
        "The EduCenter source code is organized in a modular MVC structure under project_root/. "
        "Controllers handle HTTP requests and delegate business logic to Models/Services. Views are PHP "
        "templates that display data passed from controllers. The sections below highlight key implementation "
        "areas. Red markers indicate where you should insert screenshots of the corresponding PHP source files."
    ))

    # 3.1.1
    add_heading(doc, "3.1.1. Application Entry Point and Routing", 3)
    add_para(doc, (
        "The web portal starts at public/web.php, which registers an autoloader, creates a Router instance, "
        "loads all routes from app/web_routes.php, and dispatches the request to the matching controller action."
    ))
    add_screenshot_marker(doc, "Application bootstrap and route dispatch",
                          "project_root/public/web.php",
                          "Capture lines showing session_start(), autoload, Router, and dispatch().")
    add_screenshot_marker(doc, "Custom Router class",
                          "project_root/core/router.php",
                          "Capture the dispatch() method and route matching logic.")
    add_screenshot_marker(doc, "Route definitions",
                          "project_root/app/web_routes.php",
                          "Capture admin, teacher, and student route groups with middleware.")

    # 3.1.2
    add_heading(doc, "3.1.2. Authentication and Role-Based Access Control", 3)
    add_para(doc, (
        "Authentication uses PHP sessions. WebAuthController handles login/logout. SessionAuth middleware "
        "protects routes and enforces role-based access (ADMIN, TEACHER, STUDENT)."
    ))
    add_screenshot_marker(doc, "Login controller",
                          "project_root/app/controllers/WebAuthController.php",
                          "Capture the login() method: validation, UserModel::loginWeb(), session, redirect.")
    add_screenshot_marker(doc, "Session authentication middleware",
                          "project_root/app/middleware/SessionAuth.php",
                          "Capture handle(), guest(), and requireRole() methods.")
    add_screenshot_marker(doc, "Login view template",
                          "project_root/app/view/auth/login.php",
                          "Capture the login form HTML.")

    # 3.1.3
    add_heading(doc, "3.1.3. Admin Dashboard and Payment Workflow", 3)
    add_para(doc, (
        "The admin dashboard displays key metrics. PaymentService handles tuition confirmation with "
        "schedule conflict detection before activating student enrollments."
    ))
    add_screenshot_marker(doc, "Admin dashboard controller",
                          "project_root/app/controllers/admin/AdminController.php",
                          "Capture the dashboard() method.")
    add_screenshot_marker(doc, "Admin dashboard view",
                          "project_root/app/view/admin/dashboard.php",
                          "Capture the dashboard layout with statistics cards.")
    add_screenshot_marker(doc, "Payment service — conflict check",
                          "project_root/app/models/PaymentService.php",
                          "Capture listPending() and confirmPayment() methods.")
    add_screenshot_marker(doc, "Payment management controller",
                          "project_root/app/controllers/admin/PaymentWebController.php",
                          "Capture confirm payment action.")

    # 3.1.4
    add_heading(doc, "3.1.4. Schedule Management and Conflict Detection", 3)
    add_para(doc, (
        "ScheduleService generates regular schedules, detects teacher/student conflicts, and supports "
        "week-based schedule views for teachers and students."
    ))
    add_screenshot_marker(doc, "Schedule service — conflict detection",
                          "project_root/app/models/ScheduleService.php",
                          "Capture checkTeacherConflict() and checkStudentScheduleConflict() methods.")
    add_screenshot_marker(doc, "Teacher schedule controller",
                          "project_root/app/controllers/teacher/TeacherPortalController.php",
                          "Capture schedule() method with week parameter.")
    add_screenshot_marker(doc, "Teacher schedule view with week navigation",
                          "project_root/app/view/teacher/schedule.php",
                          "Capture week date picker and Previous/Next buttons.")
    add_screenshot_marker(doc, "Student schedule controller",
                          "project_root/app/controllers/student/StudentPortalController.php",
                          "Capture schedule() method.")
    add_screenshot_marker(doc, "Student schedule view with week navigation",
                          "project_root/app/view/student/schedule.php",
                          "Capture week navigation buttons.")

    # 3.1.5
    add_heading(doc, "3.1.5. Assignment Submission and Grading", 3)
    add_para(doc, (
        "Students upload PDF assignments. Phase2Service validates enrollment, saves the file, and notifies "
        "teachers. Teachers grade submissions and the system updates student evaluations."
    ))
    add_screenshot_marker(doc, "Student submission controller",
                          "project_root/app/controllers/student/StudentPhase2Controller.php",
                          "Capture submissions() and upload/store methods.")
    add_screenshot_marker(doc, "Submission service",
                          "project_root/app/models/Phase2Service.php",
                          "Capture saveSubmission(), getSubmissionsByStudent(), gradeSubmission().")
    add_screenshot_marker(doc, "Student submissions view",
                          "project_root/app/view/student/submissions.php",
                          "Capture upload form and submission list table.")
    add_screenshot_marker(doc, "Teacher grading controller",
                          "project_root/app/controllers/teacher/TeacherPhase2Controller.php",
                          "Capture grading() and grade() methods.")
    add_screenshot_marker(doc, "Teacher grading view",
                          "project_root/app/view/teacher/grading.php",
                          "Capture grading form with score input.")

    # 3.1.6
    add_heading(doc, "3.1.6. Database Connection and Migration", 3)
    add_para(doc, (
        "Database access uses a PDO singleton. The migrate.php script creates all tables, syncs schema "
        "changes, and seeds demo data."
    ))
    add_screenshot_marker(doc, "Database singleton (PDO)",
                          "project_root/core/database.php",
                          "Capture getInstance() and PDO connection setup.")
    add_screenshot_marker(doc, "Database configuration",
                          "project_root/config/database.php",
                          "Capture host, port, dbname, username settings.")
    add_screenshot_marker(doc, "Migration script — table creation",
                          "migrate.php",
                          "Capture CREATE TABLE statements for users, courses, classes, etc.")

    # 3.1.7
    add_heading(doc, "3.1.7. MVC Separation — Controller and View Templates", 3)
    add_para(doc, (
        "Per lecturer requirements, controllers should only handle request logic and pass data to view "
        "templates (display templates). The WebController base class provides render(), flash messages, "
        "validation, and redirect helpers."
    ))
    add_screenshot_marker(doc, "Base WebController",
                          "project_root/core/WebController.php",
                          "Capture render() method showing controller → view separation.")
    add_screenshot_marker(doc, "Main layout template",
                          "project_root/app/view/layouts/main.php",
                          "Capture the shared layout with sidebar and content area.")
    add_screenshot_marker(doc, "Sidebar navigation partial",
                          "project_root/app/view/partials/sidebar.php",
                          "Capture role-based menu items.")

    add_heading(doc, "3.2 Images of Final Application", 2)
    add_para(doc, (
        "The following figures show screenshots of the running application. Access the portal at: "
        "http://localhost/educationcenterapi/project_root/public/web.php/"
    ))
    add_para(doc, "Demo accounts:", bold=True)
    add_bullets(doc, [
        "Admin: admin@edu.vn / admin123",
        "Teacher: tuan.gv@edu.vn / teacher123",
        "Student: an.hv@edu.vn / student123",
    ])

    ui_screens = [
        ("Figure 3.1: Admin Dashboard", "/admin/dashboard", "project_root/app/view/admin/dashboard.php"),
        ("Figure 3.2: Login Page", "/login", "project_root/app/view/auth/login.php"),
        ("Figure 3.3: Course Management (Admin)", "/admin/courses", "project_root/app/view/admin/courses/index.php"),
        ("Figure 3.4: Class & Enrollment Management", "/admin/classes", "project_root/app/view/admin/classes/index.php"),
        ("Figure 3.5: Teacher Schedule Page", "/teacher/schedule", "project_root/app/view/teacher/schedule.php"),
        ("Figure 3.6: Student Schedule Page", "/student/schedule", "project_root/app/view/student/schedule.php"),
        ("Figure 3.7: Student PDF Submission Page", "/student/submissions", "project_root/app/view/student/submissions.php"),
        ("Figure 3.8: Teacher Grading Page", "/teacher/grading", "project_root/app/view/teacher/grading.php"),
        ("Figure 3.9: Payment Management (Admin)", "/admin/payments", "project_root/app/view/admin/payments/index.php"),
        ("Figure 3.10: Leave/Makeup Approval (Admin)", "/admin/leave-requests", "project_root/app/view/admin/leave_requests.php"),
        ("Figure 3.11: Student Results Page", "/student/results", "project_root/app/view/student/results.php"),
        ("Figure 3.12: Attendance Report (Admin)", "/admin/attendance-report", "project_root/app/view/admin/attendance_report.php"),
    ]
    for caption, url, view_file in ui_screens:
        add_figure_placeholder(doc, f"{caption}\nURL: .../web.php{url}\nView file: {view_file}")

    add_heading(doc, "3.3 GitHub Repository Evidences", 2)
    add_para(doc, (
        "Per individual assessment requirements, each student must have at least 10 meaningful commits "
        "across at least 3 different days. Insert screenshots of the GitHub repository commit history below."
    ))
    add_bullets(doc, [
        "Repository URL: [INSERT YOUR GITHUB REPO URL HERE]",
        "Branch: main / develop",
        "Requirement: ≥10 meaningful commits per student, across ≥3 different days",
        "Commits should show work on controllers, views, models, or bug fixes",
    ])
    for i in range(1, 5):
        add_figure_placeholder(doc, f"Figure 3.{9+i}: GitHub commit history — Group Member {i}")

    doc.add_page_break()

    # ============================================================
    # CHAPTER 4
    # ============================================================
    add_heading(doc, "CHAPTER 4: CONCLUSION", 1)

    add_heading(doc, "4.1 What Went Well", 2)
    add_bullets(doc, [
        "Successfully implemented a full MVC web application with three role-based portals (Admin, Teacher, Student).",
        "Designed and deployed a normalized database with 24 tables, migration script, and seed data.",
        "Built core business workflows: enrollment → payment → schedule activation with conflict detection.",
        "Implemented schedule week navigation (Previous/Next) for both teacher and student portals.",
        "Developed PDF submission and grading workflow with notifications and student evaluations.",
        "Applied session-based authentication with role middleware protecting all sensitive routes.",
        "Created admin reporting, audit logs, CSV export, and notification system.",
        "Controllers correctly separate business logic from view templates following MVC principles.",
    ])

    add_heading(doc, "4.2 What Did Not Go Well", 2)
    add_bullets(doc, [
        "Database schema migration issues occurred when upgrading from an older schema (e.g., missing columns like total_sessions, grades.score) — resolved by ALTER sync in migrate.php.",
        "MySQL port conflict between XAMPP MySQL (port 3307) and standalone MySQL Server (port 3306) caused connection errors during deployment.",
        "Some legacy API routes (JWT-based) and new web portal routes coexist, creating maintenance complexity.",
        "PDF file upload validation and storage security could be further strengthened (file type verification, size limits).",
        "The UI design is functional but not fully responsive on all mobile screen sizes.",
        "Git version control was not initialized from the start of the project, making commit history harder to organize.",
    ])

    add_heading(doc, "4.3 Lessons Learned and Further Improvements", 2)
    add_para(doc, "Lessons learned:", bold=True)
    add_bullets(doc, [
        "Plan the database schema carefully before implementation; use migration scripts with ALTER sync for evolving schemas.",
        "Separate concerns strictly: controllers handle HTTP, models handle data, views handle display.",
        "Role-based access control should be enforced at the middleware level, not just in views.",
        "Schedule conflict detection should be integrated into payment and enrollment workflows, not checked manually.",
        "Test on the target deployment environment early (XAMPP, MySQL port, PHP version).",
        "Use Git from day one with meaningful commit messages across multiple days per team member.",
    ])
    add_para(doc, "Further improvements:", bold=True)
    add_bullets(doc, [
        "Integrate online payment gateway (VNPay, MoMo) for automatic tuition confirmation.",
        "Add email/SMS notifications in addition to in-app notifications.",
        "Improve mobile-responsive UI with a modern CSS framework (Bootstrap 5 or Tailwind).",
        "Add real-time chat between teachers and students.",
        "Implement automatic schedule generation based on course templates and room availability.",
        "Add data visualization charts (Chart.js) to admin dashboard for revenue and enrollment trends.",
        "Deploy to a cloud server with HTTPS, automated backups, and CI/CD pipeline.",
        "Write comprehensive unit and integration tests (PHPUnit) for critical business logic.",
    ])

    # ===== ASSESSMENT CRITERIA SUMMARY =====
    doc.add_page_break()
    add_heading(doc, "APPENDIX: Assessment Criteria Mapping", 1)
    add_para(doc, "Group Report (30%):", bold=True)
    add_bullets(doc, [
        "User Stories (2 pts) — Section 2.1",
        "Site Map (2 pts) — Section 2.2",
        "ERD (3 pts) — Section 2.3",
        "Final Result with evidence (3 pts) — Sections 3.2 and 3.3",
    ])
    add_para(doc, "Individual Assessment (70%):", bold=True)
    add_bullets(doc, [
        "Functionality and Presentation (3 pts) — Sections 3.1 and 3.2",
        "Technical questions (3 pts) — understand code in Section 3.1",
        "≥10 commits in ≥3 days per student (2 pts) — Section 3.3",
        "Controllers and views correctly separated (2 pts) — Section 3.1.7",
    ])

    doc.save(OUTPUT)
    print(f"Report saved to: {OUTPUT}")


if __name__ == "__main__":
    build_report()

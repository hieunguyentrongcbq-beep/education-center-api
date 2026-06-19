# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "Section_2_3_ERD_Physical_Design_v2.docx")

# ---- Physical design data: (table_title, [ (attr, type, constraint, meaning_vi) ]) ----
TABLES = [
    ("Bảng users (Người dùng)", [
        ("id", "INT", "PK, AUTO_INCREMENT", "Mã tài khoản nội bộ (duy nhất)"),
        ("full_name", "VARCHAR(150)", "NOT NULL", "Họ và tên đầy đủ"),
        ("email", "VARCHAR(150)", "UNIQUE, NOT NULL", "Email đăng nhập duy nhất"),
        ("password_hash", "VARCHAR(255)", "NOT NULL", "Mật khẩu đã mã hóa (bcrypt)"),
        ("phone", "VARCHAR(20)", "NULL", "Số điện thoại liên hệ"),
        ("status", "ENUM", "ACTIVE / INACTIVE", "Trạng thái tài khoản"),
        ("created_at", "TIMESTAMP", "DEFAULT NOW", "Thời điểm tạo tài khoản"),
        ("updated_at", "TIMESTAMP", "ON UPDATE NOW", "Lần cập nhật gần nhất"),
    ]),
    ("Bảng roles (Vai trò) & user_roles", [
        ("roles.id", "INT", "PK", "Mã vai trò"),
        ("roles.role_name", "VARCHAR(60)", "UNIQUE", "Tên vai trò: ADMIN / TEACHER / STUDENT"),
        ("roles.description", "TEXT", "NULL", "Mô tả vai trò"),
        ("user_roles.user_id", "INT", "FK -> users", "Tham chiếu người dùng"),
        ("user_roles.role_id", "INT", "FK -> roles", "Tham chiếu vai trò"),
    ]),
    ("Bảng teachers (Giáo viên)", [
        ("id", "INT", "PK", "Mã giáo viên nội bộ"),
        ("user_id", "INT", "FK -> users", "Tài khoản liên kết"),
        ("teacher_code", "VARCHAR(30)", "UNIQUE", "Mã giáo viên hiển thị"),
        ("specialization", "VARCHAR(200)", "NOT NULL", "Chuyên môn giảng dạy"),
        ("hire_date", "DATE", "NULL", "Ngày vào làm"),
        ("teacher_type", "ENUM", "FULL_TIME / VISITING", "Loại giáo viên"),
        ("standard_hours", "DECIMAL(5,1)", "DEFAULT 40", "Số giờ chuẩn"),
        ("status", "ENUM", "ACTIVE / INACTIVE", "Trạng thái"),
    ]),
    ("Bảng students (Học viên)", [
        ("id", "INT", "PK", "Mã học viên nội bộ"),
        ("user_id", "INT", "FK -> users", "Tài khoản liên kết"),
        ("student_code", "VARCHAR(30)", "UNIQUE", "Mã học viên hiển thị"),
        ("date_of_birth", "DATE", "NULL", "Ngày sinh"),
        ("parent_phone", "VARCHAR(20)", "NULL", "SĐT phụ huynh"),
        ("status", "ENUM", "ACTIVE / INACTIVE / GRADUATED", "Trạng thái học viên"),
    ]),
    ("Bảng courses (Khóa học)", [
        ("id", "INT", "PK", "Mã khóa học nội bộ"),
        ("course_code", "VARCHAR(30)", "UNIQUE", "Mã khóa học (VD: JS101)"),
        ("course_name", "VARCHAR(200)", "NOT NULL", "Tên khóa học"),
        ("description", "TEXT", "NULL", "Mô tả khóa học"),
        ("duration_weeks", "INT", "DEFAULT 10", "Số tuần học"),
        ("total_sessions", "INT", "DEFAULT 20", "Tổng số buổi học"),
        ("day_primary", "TINYINT", "DEFAULT 1", "Thứ học chính"),
        ("day_secondary", "TINYINT", "DEFAULT 4", "Thứ học phụ"),
        ("tuition_fee", "DECIMAL(12,2)", "NOT NULL", "Học phí"),
        ("status", "ENUM", "ACTIVE / INACTIVE", "Trạng thái khóa học"),
    ]),
    ("Bảng classrooms (Phòng học) & semesters (Học kỳ)", [
        ("classrooms.id", "INT", "PK", "Mã phòng học"),
        ("classrooms.room_name", "VARCHAR(50)", "UNIQUE", "Tên phòng"),
        ("classrooms.capacity", "INT", "NOT NULL", "Sức chứa"),
        ("semesters.id", "INT", "PK", "Mã học kỳ"),
        ("semesters.semester_name", "VARCHAR(100)", "NOT NULL", "Tên học kỳ"),
        ("semesters.start_date", "DATE", "NOT NULL", "Ngày bắt đầu"),
        ("semesters.end_date", "DATE", "NOT NULL", "Ngày kết thúc"),
    ]),
    ("Bảng classes (Lớp học)", [
        ("id", "INT", "PK", "Mã lớp nội bộ"),
        ("class_code", "VARCHAR(40)", "UNIQUE", "Mã lớp (VD: JS101-A)"),
        ("course_id", "INT", "FK -> courses", "Khóa học của lớp"),
        ("teacher_id", "INT", "FK -> teachers", "Giáo viên phụ trách"),
        ("classroom_id", "INT", "FK -> classrooms", "Phòng học"),
        ("semester_id", "INT", "FK -> semesters", "Học kỳ"),
        ("max_students", "INT", "DEFAULT 30", "Sĩ số tối đa"),
        ("start_date", "DATE", "NOT NULL", "Ngày khai giảng"),
        ("end_date", "DATE", "NOT NULL", "Ngày kết thúc"),
        ("status", "ENUM", "UPCOMING / ONGOING / COMPLETED / CANCELLED", "Trạng thái lớp"),
    ]),
    ("Bảng enrollments (Ghi danh)", [
        ("id", "INT", "PK", "Mã ghi danh"),
        ("student_id", "INT", "FK -> students", "Học viên"),
        ("class_id", "INT", "FK -> classes", "Lớp học"),
        ("enrollment_date", "DATE", "NOT NULL", "Ngày ghi danh"),
        ("payment_status", "ENUM", "UNPAID / PAID / REFUNDED", "Trạng thái thanh toán"),
        ("status", "ENUM", "ACTIVE / DROPPED / COMPLETED", "Trạng thái học"),
        ("(student_id, class_id)", "UNIQUE", "UNIQUE KEY", "Chống ghi danh trùng"),
    ]),
    ("Bảng schedules (Lịch học)", [
        ("id", "INT", "PK", "Mã lịch"),
        ("class_id", "INT", "FK -> classes", "Lớp học"),
        ("student_id", "INT", "FK -> students, NULL", "Lịch riêng theo học viên"),
        ("day_of_week", "TINYINT", "0-6", "Thứ trong tuần"),
        ("specific_date", "DATE", "NULL", "Ngày cụ thể (thi/học bù)"),
        ("start_time", "TIME", "NOT NULL", "Giờ bắt đầu"),
        ("end_time", "TIME", "NOT NULL", "Giờ kết thúc"),
        ("schedule_type", "ENUM", "REGULAR / EXAM / MAKEUP / EXTRA", "Loại buổi học"),
    ]),
    ("Bảng teacher_assignments (Phân công giáo viên)", [
        ("id", "INT", "PK", "Mã phân công"),
        ("teacher_id", "INT", "FK -> teachers", "Giáo viên"),
        ("class_id", "INT", "FK -> classes", "Lớp học"),
        ("student_id", "INT", "FK -> students, NULL", "Học viên (phân công 1-1)"),
        ("day_of_week", "TINYINT", "NULL", "Thứ dạy"),
        ("assignment_status", "ENUM", "PENDING / CONFIRMED / CANCELLED", "Trạng thái phân công"),
    ]),
    ("Bảng attendance (Điểm danh)", [
        ("id", "INT", "PK", "Mã điểm danh"),
        ("class_id", "INT", "FK -> classes", "Lớp học"),
        ("student_id", "INT", "FK -> students", "Học viên"),
        ("attendance_date", "DATE", "NOT NULL", "Ngày điểm danh"),
        ("attendance_status", "ENUM", "PRESENT / ABSENT / LATE / EXCUSED", "Trạng thái điểm danh"),
        ("tinh_luong", "TINYINT(1)", "DEFAULT 0", "Cờ tính lương giáo viên"),
        ("note", "VARCHAR(300)", "NULL", "Ghi chú"),
    ]),
    ("Bảng submissions (Bài nộp) & grades (Điểm)", [
        ("submissions.id", "INT", "PK", "Mã bài nộp"),
        ("submissions.student_id", "INT", "FK -> students", "Học viên nộp"),
        ("submissions.class_id", "INT", "FK -> classes", "Lớp học"),
        ("submissions.type", "ENUM", "ASSIGNMENT / MIDTERM / FINAL", "Loại bài"),
        ("submissions.file_path", "VARCHAR(500)", "NULL", "Đường dẫn file PDF"),
        ("submissions.status", "ENUM", "PENDING / GRADED", "Trạng thái chấm"),
        ("grades.submission_id", "INT", "FK -> submissions", "Bài nộp được chấm"),
        ("grades.score", "DECIMAL(5,2)", "NULL", "Điểm số (0-10)"),
        ("grades.comment", "TEXT", "NULL", "Nhận xét của giáo viên"),
    ]),
    ("Bảng tuition_payments (Thanh toán học phí) & payrolls (Lương GV)", [
        ("tuition_payments.id", "INT", "PK", "Mã thanh toán"),
        ("tuition_payments.student_id", "INT", "FK -> students", "Học viên"),
        ("tuition_payments.class_id", "INT", "FK -> classes, NULL", "Lớp học"),
        ("tuition_payments.amount", "DECIMAL(12,2)", "NOT NULL", "Số tiền"),
        ("tuition_payments.payment_status", "ENUM", "UNPAID / COMPLETED / REFUNDED", "Trạng thái"),
        ("payrolls.teacher_id", "INT", "FK -> teachers", "Giáo viên"),
        ("payrolls.month", "VARCHAR(7)", "NOT NULL", "Tháng tính lương"),
        ("payrolls.salary_amount", "DECIMAL(12,2)", "DEFAULT 0", "Lương"),
        ("payrolls.payment_status", "ENUM", "PENDING / PAID", "Trạng thái trả lương"),
    ]),
    ("Bảng surveys (Khảo sát) & survey_responses (Phản hồi)", [
        ("surveys.id", "INT", "PK", "Mã khảo sát"),
        ("surveys.class_id", "INT", "FK -> classes", "Lớp khảo sát"),
        ("surveys.title", "VARCHAR(200)", "NOT NULL", "Tiêu đề khảo sát"),
        ("survey_responses.survey_id", "INT", "FK -> surveys", "Khảo sát"),
        ("survey_responses.student_id", "INT", "FK -> students", "Học viên đánh giá"),
        ("survey_responses.teacher_id", "INT", "FK -> teachers", "Giáo viên được đánh giá"),
        ("survey_responses.rating", "TINYINT", "NOT NULL", "Điểm đánh giá"),
    ]),
    ("Bảng notifications (Thông báo) & audit_logs (Nhật ký)", [
        ("notifications.id", "INT", "PK", "Mã thông báo"),
        ("notifications.title", "VARCHAR(200)", "NOT NULL", "Tiêu đề"),
        ("notifications.receiver_id", "INT", "FK -> users", "Người nhận"),
        ("notifications.is_read", "TINYINT(1)", "DEFAULT 0", "Đã đọc hay chưa"),
        ("audit_logs.user_id", "INT", "FK -> users", "Người thực hiện"),
        ("audit_logs.action", "VARCHAR(60)", "NULL", "Hành động (CREATE/UPDATE/DELETE)"),
        ("audit_logs.entity_name", "VARCHAR(60)", "NULL", "Bảng bị tác động"),
        ("audit_logs.created_at", "TIMESTAMP", "DEFAULT NOW", "Thời điểm ghi log"),
    ]),
]

# ---- Key relationships: (A, rel, B, desc_vi) ----
RELATIONS = [
    ("users", "1:1", "teachers", "Một tài khoản liên kết một hồ sơ giáo viên"),
    ("users", "1:1", "students", "Một tài khoản liên kết một hồ sơ học viên"),
    ("users", "M:N", "roles", "Người dùng và vai trò qua bảng user_roles"),
    ("courses", "1:N", "classes", "Một khóa học có nhiều lớp"),
    ("semesters", "1:N", "classes", "Một học kỳ chứa nhiều lớp"),
    ("classrooms", "1:N", "classes", "Một phòng học phục vụ nhiều lớp"),
    ("classes", "1:N", "schedules", "Một lớp có nhiều buổi/lịch học"),
    ("classes", "1:N", "enrollments", "Một lớp có nhiều lượt ghi danh"),
    ("students", "1:N", "enrollments", "Một học viên ghi danh nhiều lớp"),
    ("enrollments", "1:N", "tuition_payments", "Một ghi danh có thể có nhiều thanh toán"),
    ("classes", "1:N", "teacher_assignments", "Một lớp có các phân công giáo viên"),
    ("teachers", "1:N", "teacher_assignments", "Một giáo viên được phân công nhiều lớp"),
    ("classes", "1:N", "attendance", "Một lớp có nhiều bản ghi điểm danh"),
    ("students", "1:N", "attendance", "Một học viên có nhiều bản ghi điểm danh"),
    ("students", "1:N", "submissions", "Một học viên nộp nhiều bài"),
    ("submissions", "1:1", "grades", "Mỗi bài nộp có một điểm"),
    ("teachers", "1:N", "payrolls", "Một giáo viên có nhiều bản ghi lương"),
    ("classes", "1:1", "surveys", "Một lớp có một khảo sát"),
    ("surveys", "1:N", "survey_responses", "Một khảo sát có nhiều phản hồi"),
    ("users", "1:N", "notifications", "Một người dùng nhận nhiều thông báo"),
    ("users", "1:N", "audit_logs", "Một người dùng sinh nhiều dòng nhật ký"),
]

RELATIONS_EN = [
    ("users", "1:1", "teachers", "One user account links to one teacher profile"),
    ("users", "1:1", "students", "One user account links to one student profile"),
    ("users", "M:N", "roles", "Users and roles are mapped through user_roles"),
    ("courses", "1:N", "classes", "One course can contain many classes"),
    ("semesters", "1:N", "classes", "One semester contains many classes"),
    ("classrooms", "1:N", "classes", "One classroom can host many classes"),
    ("classes", "1:N", "schedules", "One class has multiple schedule slots"),
    ("classes", "1:N", "enrollments", "One class has many enrollments"),
    ("students", "1:N", "enrollments", "One student can enroll in many classes"),
    ("enrollments", "1:N", "tuition_payments", "One enrollment can have multiple payment records"),
    ("classes", "1:N", "teacher_assignments", "One class has multiple teacher assignments"),
    ("teachers", "1:N", "teacher_assignments", "One teacher can be assigned to many classes"),
    ("classes", "1:N", "attendance", "One class has many attendance records"),
    ("students", "1:N", "attendance", "One student has many attendance records"),
    ("students", "1:N", "submissions", "One student can upload multiple submissions"),
    ("submissions", "1:1", "grades", "One submission has one grading result"),
    ("teachers", "1:N", "payrolls", "One teacher can have multiple payroll records"),
    ("classes", "1:1", "surveys", "One class has one survey"),
    ("surveys", "1:N", "survey_responses", "One survey collects many responses"),
    ("users", "1:N", "notifications", "One user can receive many notifications"),
    ("users", "1:N", "audit_logs", "One user can generate many audit log entries"),
]

PHYSICAL_SUMMARY_EN = [
    ("users", "id, full_name, email, password_hash, status", "Stores system user accounts and login identity"),
    ("roles / user_roles", "role_name, user_id, role_id", "Implements role-based access control (RBAC)"),
    ("teachers / students", "teacher_code, student_code, user_id", "Extends user profiles for teacher and student domains"),
    ("courses", "course_code, course_name, total_sessions, tuition_fee", "Stores course catalog and tuition information"),
    ("classes", "class_code, course_id, teacher_id, semester_id, status", "Stores class instances opened from courses"),
    ("enrollments", "student_id, class_id, payment_status, status", "Tracks student-class registration and learning status"),
    ("schedules", "class_id, student_id, day_of_week, specific_date", "Stores regular and specific-date learning sessions"),
    ("teacher_assignments", "teacher_id, class_id, student_id, assignment_status", "Assigns teachers to classes/students"),
    ("attendance", "class_id, student_id, attendance_date, attendance_status", "Tracks daily attendance outcomes"),
    ("submissions / grades", "submission_id, score, comment", "Handles assignment submission and grading workflow"),
    ("tuition_payments / payrolls", "amount, payment_status, month, salary_amount", "Tracks tuition and payroll transactions"),
    ("surveys / survey_responses", "class_id, teacher_id, rating", "Collects end-of-course teacher feedback"),
    ("notifications / audit_logs", "receiver_id, action, entity_name, created_at", "Supports alerts and system activity traceability"),
]


def style_runs(cell, size=9, bold=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(size)
            r.bold = bold


def add_heading(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, bold=False, italic=False, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def add_design_table(doc, title, rows, index):
    add_para(doc, f"Table 2.{index}. {title}", bold=True)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Thuộc tính", "Kiểu dữ liệu", "Ràng buộc", "Ý nghĩa"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        style_runs(table.rows[0].cells[i], 9, True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            style_runs(cells[i], 9)
    doc.add_paragraph()


def build():
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    add_heading(doc, "2.3 Entity Relationship Diagram (ERD)", 2)
    add_para(doc, (
        "Cơ sở dữ liệu EduCenter sử dụng mô hình quan hệ gồm 24 bảng. Các thực thể trung tâm gồm "
        "users, roles, teachers, students, courses, classes, enrollments, schedules và teacher_assignments. "
        "Phần dưới đây trình bày thiết kế vật lý (Physical Database Design) của các bảng chính và bảng "
        "tổng hợp quan hệ giữa các thực thể."
    ))
    doc.add_paragraph()

    add_heading(doc, "2.3.1 Physical Database Design", 3)
    for i, (title, rows) in enumerate(TABLES, start=1):
        add_design_table(doc, title, rows, i)

    doc.add_page_break()

    add_heading(doc, "2.3.2 Key Relationships", 3)
    add_para(doc, "Bảng dưới đây tóm tắt các quan hệ chính giữa các thực thể trong cơ sở dữ liệu.")
    rel_table = doc.add_table(rows=1, cols=4)
    rel_table.style = "Table Grid"
    headers = ["Entity A", "Relationship", "Entity B", "Mô tả (Description)"]
    for i, h in enumerate(headers):
        rel_table.rows[0].cells[i].text = h
        style_runs(rel_table.rows[0].cells[i], 10, True)
    for a, rel, b, desc in RELATIONS:
        cells = rel_table.add_row().cells
        cells[0].text = a
        cells[1].text = rel
        cells[2].text = b
        cells[3].text = desc
        for c in cells:
            style_runs(c, 9)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[INSERT ERD IMAGE HERE — Figure 2.2: Entity Relationship Diagram — finaldemo_web Database]")
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    run.font.name = "Times New Roman"

    doc.add_page_break()
    add_heading(doc, "2.3.3 Key Relationships (English Version)", 3)
    add_para(doc, "The table below summarizes the primary relationships between entities in English.")
    rel_en = doc.add_table(rows=1, cols=4)
    rel_en.style = "Table Grid"
    headers_en = ["Entity A", "Relationship", "Entity B", "Description"]
    for i, h in enumerate(headers_en):
        rel_en.rows[0].cells[i].text = h
        style_runs(rel_en.rows[0].cells[i], 10, True)
    for a, rel, b, desc in RELATIONS_EN:
        cells = rel_en.add_row().cells
        cells[0].text = a
        cells[1].text = rel
        cells[2].text = b
        cells[3].text = desc
        for c in cells:
            style_runs(c, 9)

    doc.add_paragraph()
    add_heading(doc, "2.3.4 Physical Database Design (English Summary)", 3)
    add_para(doc, "Compact English summary of major tables for quick reference.")
    summ = doc.add_table(rows=1, cols=3)
    summ.style = "Table Grid"
    for i, h in enumerate(["Entity/Table", "Key Fields", "Purpose"]):
        summ.rows[0].cells[i].text = h
        style_runs(summ.rows[0].cells[i], 10, True)
    for entity, fields, purpose in PHYSICAL_SUMMARY_EN:
        cells = summ.add_row().cells
        cells[0].text = entity
        cells[1].text = fields
        cells[2].text = purpose
        for c in cells:
            style_runs(c, 9)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

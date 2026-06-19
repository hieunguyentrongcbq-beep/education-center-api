# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "Section_2_1_User_Stories_File_Map.docx")

ROWS = [
    # US, Criteria, File PHP
    ("US-01", "Create/edit/delete courses", "Controller: app/controllers/admin/CourseWebController.php\nModel: app/models/classmodel.php → createCourse(), updateCourse(), deleteCourse()\nView: app/view/admin/courses/index.php, form.php"),
    ("US-01", "Create/manage classes", "Controller: app/controllers/admin/ClassWebController.php\nModel: app/models/classmodel.php → createClass(), updateClass(), deleteClass()\nView: app/view/admin/classes/index.php, form.php"),
    ("US-01", "Linked to enrollments/schedules/payments", "Model: app/models/PaymentService.php, ScheduleService.php"),

    ("US-02", "Admin selects student + class", "Controller: app/controllers/admin/EnrollmentWebController.php → store()\nView: app/view/admin/enrollments/form.php"),
    ("US-02", "payment_status=UNPAID, status=ACTIVE", "Model: app/models/classmodel.php → createEnrollment() ~1636-1643"),
    ("US-02", "Prevent duplicate enrollment", "Model: app/models/classmodel.php → ~1629-1631"),
    ("US-02", "Auto create tuition payment", "Model: app/models/classmodel.php → ~1645-1655"),

    ("US-03", "View & confirm pending payments", "Controller: app/controllers/admin/PaymentWebController.php\nModel: app/models/PaymentService.php → listPending(), confirmPayment()\nView: app/view/admin/payments/index.php"),
    ("US-03", "Update payment_status=PAID", "Model: app/models/PaymentService.php → ~250-264"),
    ("US-03", "Activate student schedule", "Model: PaymentService.php → ~272; ScheduleService.php → generateRegularSchedule()"),
    ("US-03", "Notify student", "Model: PaymentService.php → ~288-296"),

    ("US-04", "Check schedule overlap", "Model: app/models/ScheduleService.php → checkStudentScheduleConflict()"),
    ("US-04", "Block confirmation + error", "Model: app/models/PaymentService.php → ~236-239"),
    ("US-04", "Show conflict on pending list", "Model: PaymentService.php → listPending() ~44-47\nView: app/view/admin/payments/index.php"),

    ("US-05", "Assign teacher to class/student", "Controller: app/controllers/admin/AssignmentWebController.php → store()\nModel: app/models/assigning.php → assignTeacher()\nView: app/view/admin/assignments/form.php"),
    ("US-05", "Prevent teacher double-booking", "Model: assigning.php → checkTeacherConflict\nModel: ScheduleService.php → checkTeacherConflict()"),
    ("US-05", "Visible on admin + teacher schedule", "View: app/view/admin/assignments/teaching_schedule.php\nController: app/controllers/teacher/TeacherPortalController.php → schedule()\nView: app/view/teacher/schedule.php"),

    ("US-06", "Dashboard stats", "Controller: app/controllers/admin/AdminController.php → dashboard()\nModel: classmodel.php → getRevenueByMonth(), getStudentCount(), getClassFillRate()\nView: app/view/admin/dashboard.php"),
    ("US-06", "Pending payment alerts", "Model: PaymentService.php → listPending(), countPending()\nView: app/view/admin/dashboard.php"),

    ("US-07", "List leave/makeup requests", "Controller: app/controllers/admin/AdminPhase2Controller.php → leaveRequests()\nModel: Phase2Service.php → listLeaveRequests()\nView: app/view/admin/leave_requests.php"),
    ("US-07", "Approve/reject", "Controller: AdminPhase2Controller.php → reviewLeave()\nModel: Phase2Service.php → reviewLeaveRequest()"),
    ("US-07", "MAKEUP → create session", "Model: Phase2Service.php → createMakeupScheduleFromLeave()"),
    ("US-07", "Notify requester", "Model: Phase2Service.php → notifyLeaveReviewed()"),

    ("US-08", "Export revenue CSV", "Controller: app/controllers/admin/AdminController.php → exportRevenue()"),
    ("US-08", "View/filter audit logs", "Controller: AdminPhase2Controller.php → auditLogs()\nModel: Phase2Service.php → getAuditLogs()\nView: app/view/admin/audit_logs.php"),
    ("US-08", "Export audit CSV", "Controller: AdminPhase2Controller.php → exportAuditLogs()"),
    ("US-08", "Record audit actions", "Model: app/models/AuditLog.php → write()"),

    ("US-09", "Create/update payroll", "Controller: app/controllers/admin/PayrollWebController.php\nModel: app/models/instructormodel.php\nView: app/view/admin/payrolls/form.php, index.php"),
    ("US-09", "Mark payroll as paid", "Controller: PayrollWebController.php → markPaid()"),

    ("US-10", "Weekly teaching schedule", "Controller: app/controllers/teacher/TeacherPortalController.php → schedule()\nModel: ScheduleService.php → getTeacherSchedule()\nView: app/view/teacher/schedule.php"),
    ("US-10", "Link to attendance", "View: app/view/teacher/schedule.php → attendance link ~5-62"),

    ("US-11", "Previous/Next week buttons", "View: app/view/teacher/schedule.php ~3-41"),
    ("US-11", "Week parameter ?week=", "Controller: TeacherPortalController.php → schedule()\nModel: ScheduleService.php → normalizeWeekStart()"),

    ("US-12", "Open attendance form", "Controller: app/controllers/teacher/TeacherPhase2Controller.php → attendance()\nView: app/view/teacher/attendance.php"),
    ("US-12", "Mark PRESENT/ABSENT/LATE/EXCUSED", "Controller: TeacherPhase2Controller.php → markAttendance()\nModel: classmodel.php → markAttendance()"),
    ("US-12", "Admin attendance report", "Controller: AdminPhase2Controller.php → attendanceReport()\nView: app/view/admin/attendance_report.php"),

    ("US-13", "List submissions", "Controller: TeacherPhase2Controller.php → grading()\nModel: Phase2Service.php → getSubmissionsForTeacher()\nView: app/view/teacher/grading.php"),
    ("US-13", "Grade score + comment", "Controller: TeacherPhase2Controller.php → grade()\nModel: Phase2Service.php → gradeSubmission()"),

    ("US-14", "Submit leave/makeup", "Controller: TeacherPhase2Controller.php → leave(), submitLeave()\nModel: Phase2Service.php → createLeaveRequest()\nView: app/view/teacher/leave.php"),

    ("US-15", "Notify on new submission", "Model: Phase2Service.php → saveSubmission(), notifyTeachersOfClass()"),
    ("US-15", "Mark notification read", "Controller: NotificationWebController.php\nJS: public/assets/js/notifications-ajax.js"),

    ("US-16", "Student weekly schedule", "Controller: app/controllers/student/StudentPortalController.php → schedule()\nModel: ScheduleService.php → getStudentScheduleForWeek()\nView: app/view/student/schedule.php"),

    ("US-17", "Previous/Next week", "View: app/view/student/schedule.php ~5-68\nController: StudentPortalController.php → schedule()"),

    ("US-18", "View attendance list", "Controller: StudentPhase2Controller.php → attendance()\nModel: Phase2Service.php → getStudentAttendanceStatus()\nView: app/view/student/attendance.php"),
    ("US-18", "Attendance on schedule", "View: app/view/student/schedule.php ~116-120"),

    ("US-19", "Upload PDF form", "View: app/view/student/submissions.php ~13-37"),
    ("US-19", "PDF max 5MB validation", "Core: core/WebController.php → uploadPdf() ~176"),
    ("US-19", "Save + notify teacher", "Controller: StudentPhase2Controller.php → upload()\nModel: Phase2Service.php → saveSubmission()"),
    ("US-19", "Download past submissions", "Controller: StudentPhase2Controller.php → downloadSubmission()\nView: submissions.php"),

    ("US-20", "View grades & comments", "Controller: StudentPhase2Controller.php → results()\nModel: Phase2Service.php → getEvaluation(), getSubmissionsByStudent()\nView: app/view/student/results.php"),

    ("US-21", "Class ranking", "Controller: StudentPhase2Controller.php → compare()\nModel: Phase2Service.php → getClassScoreComparison()\nView: app/view/student/compare.php"),

    ("US-22", "Submit leave/makeup", "Controller: StudentPhase2Controller.php → leave(), submitLeave()\nModel: Phase2Service.php → createLeaveRequest()\nView: app/view/student/leave.php"),

    ("US-23", "Survey when class COMPLETED", "Model: Phase2Service.php → getSurveyableClasses()"),
    ("US-23", "Submit survey", "Controller: StudentPhase2Controller.php → survey(), submitSurvey()\nModel: Phase2Service.php → submitSurvey()\nView: app/view/student/survey.php"),
]


def build():
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    h = doc.add_heading("2.1 User Stories — File PHP Mapping", level=2)
    for r in h.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "Base path: project_root/ — Bảng map từng chức năng User Story tới file PHP cần chụp code."
    )
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["User Story", "Chức năng / Criteria", "File PHP (chụp code)"]
    for i, text in enumerate(headers):
        hdr[i].text = text
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)

    for us, criteria, files in ROWS:
        row = table.add_row().cells
        row[0].text = us
        row[1].text = criteria
        row[2].text = files
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(9)

    table.columns[0].width = Inches(0.7)
    table.columns[1].width = Inches(1.8)
    table.columns[2].width = Inches(4.0)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

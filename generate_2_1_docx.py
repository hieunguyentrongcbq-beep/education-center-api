# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "Section_2_1_User_Stories.docx")


def normal(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)


def us_block(doc, code, story, points):
    normal(doc, f"{code}: {story}", bold=True)
    for pt in points:
        bullet(doc, pt)
    doc.add_paragraph()


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    h = doc.add_heading("2.1 User Stories", level=2)
    for r in h.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)

    normal(doc, (
        'The following user stories were produced using the standard User Story template: '
        '"As a [role], I want [feature], so that [benefit]."'
    ))
    doc.add_paragraph()

    doc.add_heading("Administrator Stories", level=3)
    us_block(doc, "US-01",
        "As an Admin, I want to manage courses and classes, so that I can organize the center's academic offerings.",
        [
            "Create, edit, and delete courses (code, name, description, duration, sessions, fee).",
            "Create and manage classes linked to course, semester, and classroom.",
        ])
    us_block(doc, "US-02",
        "As an Admin, I want to enroll students into classes, so that each student is registered for the correct program.",
        [
            "Assign a student to an open class; default payment_status = UNPAID, status = ACTIVE.",
            "Prevent duplicate enrollment; auto-create tuition payment record.",
            "Code: project_root/app/models/classmodel.php → createEnrollment()",
        ])
    us_block(doc, "US-03",
        "As an Admin, I want to confirm tuition payments, so that students can join the class and receive schedules.",
        [
            "View and confirm pending payments; update status to PAID.",
            "Activate student schedule and send in-app notification after confirmation.",
        ])
    us_block(doc, "US-04",
        "As an Admin, I want the system to detect schedule conflicts before confirming payment, so that no student is assigned to overlapping classes.",
        [
            "Check day/time overlap across enrolled classes before activation.",
            "Block confirmation and show a clear error if conflict exists.",
        ])
    us_block(doc, "US-05",
        "As an Admin, I want to assign teachers to classes and students, so that teaching responsibilities are clearly defined.",
        [
            "Assign teachers from the assignment module.",
            "Prevent teacher double-booking; assignments appear on teacher schedule.",
        ])
    us_block(doc, "US-06",
        "As an Admin, I want to view a dashboard with revenue and enrollment statistics, so that I can monitor business performance.",
        [
            "Show monthly revenue, student status, class fill rates, and pending payments.",
            "Data updates when the page is reloaded after new transactions.",
        ])
    us_block(doc, "US-07",
        "As an Admin, I want to approve or reject leave and makeup requests, so that schedule changes are controlled.",
        [
            "Review pending requests from teachers and students.",
            "Approved MAKEUP requests auto-create a makeup session; requester is notified.",
        ])
    us_block(doc, "US-08",
        "As an Admin, I want to export reports and audit logs, so that I can review system activity.",
        [
            "Export revenue and audit logs as CSV.",
            "Filter audit logs by action, entity, and date; each entry records user and timestamp.",
        ])
    us_block(doc, "US-09",
        "As an Admin, I want to manage teacher payroll, so that salary records are tracked correctly.",
        [
            "Create/update monthly payroll records and mark them as paid.",
        ])

    doc.add_heading("Teacher Stories", level=3)
    us_block(doc, "US-10",
        "As a Teacher, I want to view my weekly teaching schedule, so that I know when and where to teach.",
        [
            "View sessions by day/time for the selected week (class, date, time, student).",
            "Open attendance directly from a session.",
        ])
    us_block(doc, "US-11",
        "As a Teacher, I want to navigate between weeks (Previous/Next), so that I can plan ahead easily.",
        [
            "Move between weeks via ?week=YYYY-MM-DD; selected week shows date range.",
        ])
    us_block(doc, "US-12",
        "As a Teacher, I want to mark student attendance for each session, so that attendance records are accurate.",
        [
            "Select class and session date; mark PRESENT, ABSENT, LATE, or EXCUSED.",
            "Records are saved and visible in admin attendance report.",
        ])
    us_block(doc, "US-13",
        "As a Teacher, I want to grade student PDF submissions, so that students receive timely feedback.",
        [
            "View/download submissions; enter score (0–10) and comment visible to student.",
        ])
    us_block(doc, "US-14",
        "As a Teacher, I want to submit leave and makeup requests, so that schedule changes are formally recorded.",
        [
            "Submit LEAVE/MAKEUP request with class, date, and reason; admin reviews and notifies teacher.",
        ])
    us_block(doc, "US-15",
        "As a Teacher, I want to receive notifications when students submit assignments, so that I can grade promptly.",
        [
            "Auto notification on new submission; teacher can mark as read.",
        ])

    doc.add_heading("Student Stories", level=3)
    us_block(doc, "US-16",
        "As a Student, I want to view my class schedule by week, so that I know when to attend classes.",
        [
            "View paid, active enrollments by day/time (class, course, teacher).",
        ])
    us_block(doc, "US-17",
        "As a Student, I want to navigate between weeks (Previous/Next), so that I can check future sessions.",
        [
            "Previous/Next buttons with week date range via ?week=YYYY-MM-DD.",
        ])
    us_block(doc, "US-18",
        "As a Student, I want to view my attendance status, so that I know whether I was marked present or absent.",
        [
            "List records by class and date with PRESENT, ABSENT, LATE, or EXCUSED.",
        ])
    us_block(doc, "US-19",
        "As a Student, I want to upload PDF assignments, so that I can submit homework online.",
        [
            "Upload PDF (max 5MB) per class; view/download past submissions; teacher is notified.",
        ])
    us_block(doc, "US-20",
        "As a Student, I want to view my grades and teacher comments, so that I can track my academic progress.",
        [
            "View scores and comments after teacher grading on the results page.",
        ])
    us_block(doc, "US-21",
        "As a Student, I want to compare my scores with classmates, so that I understand my relative performance.",
        [
            "Class ranking with names, average scores, and levels; own row highlighted.",
        ])
    us_block(doc, "US-22",
        "As a Student, I want to submit leave and makeup requests, so that absences are officially recorded.",
        [
            "Submit request with class, date, and reason; admin approves/rejects and notifies student.",
        ])
    us_block(doc, "US-23",
        "As a Student, I want to complete a teacher survey after the course ends, so that I can provide feedback.",
        [
            "Survey available when class is COMPLETED; one response per student per class.",
        ])

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
